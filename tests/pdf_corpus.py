"""
Builds small, real PDFs carrying specific structural features.

These are genuine PDF files with genuine object graphs, not fixtures that
mock a parser. Every security test opens one of these through pypdf exactly
the way the application opens an upload, so a passing test means the analyzer
actually found the construct in a real file.
"""
import io
from typing import Dict, List, Optional

CONTENT_STREAM = b"BT /F1 12 Tf 72 720 Td (Sample document text. Second sentence.) Tj ET"


def build_pdf(
    catalog_extra: str = "",
    page_extra: str = "",
    extra_objects: Optional[List[bytes]] = None,
    info_extra: str = "",
    pages: int = 1,
    content: bytes = CONTENT_STREAM,
) -> bytes:
    """Assemble a valid single- or multi-page PDF.

    catalog_extra / page_extra are raw dictionary fragments spliced into the
    catalog and first page. extra_objects are appended and numbered from the
    first free object number, which callers reference as "N 0 R".
    """
    extra_objects = extra_objects or []

    kids = " ".join(f"{3 + i} 0 R" for i in range(pages))
    objects: List[bytes] = [
        f"<< /Type /Catalog /Pages 2 0 R {catalog_extra} >>".encode(),
        f"<< /Type /Pages /Kids [{kids}] /Count {pages} >>".encode(),
    ]

    # Page objects occupy 3 .. 3+pages-1
    content_obj = 3 + pages
    font_obj = content_obj + 1
    for i in range(pages):
        extras = page_extra if i == 0 else ""
        objects.append(
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            f"/Resources << /Font << /F1 {font_obj} 0 R >> >> "
            f"/Contents {content_obj} 0 R {extras} >>".encode()
        )

    objects.append(b"<< /Length %d >>\nstream\n" % len(content) + content + b"\nendstream")
    objects.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    info_obj = font_obj + 1
    objects.append(
        f"<< /Title (Corpus Document) /Author (Test Harness) {info_extra} >>".encode()
    )

    objects.extend(extra_objects)

    out = bytearray(b"%PDF-1.7\n")
    offsets: List[int] = []
    for i, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += b"%d 0 obj\n" % i + body + b"\nendobj\n"

    xref = len(out)
    out += b"xref\n0 %d\n0000000000 65535 f \n" % (len(objects) + 1)
    for off in offsets:
        out += b"%010d 00000 n \n" % off
    out += (b"trailer\n<< /Size %d /Root 1 0 R /Info %d 0 R >>\nstartxref\n%d\n%%%%EOF\n"
            % (len(objects) + 1, info_obj, xref))

    return bytes(out)


def first_free_object(pages: int = 1) -> int:
    """Object number of the first caller-supplied extra object."""
    # catalog, pages, N page objects, content, font, info
    return 2 + pages + 3 + 1


# --------------------------------------------------------------- specimens

def clean_pdf() -> bytes:
    """A plain document with no active content."""
    return build_pdf()


def multipage_pdf(pages: int = 3) -> bytes:
    return build_pdf(pages=pages)


def javascript_pdf() -> bytes:
    """JavaScript reachable through the document name tree."""
    n = first_free_object()
    return build_pdf(
        catalog_extra=f"/Names << /JavaScript {n} 0 R >>",
        extra_objects=[
            f"<< /Names [(evil) {n + 1} 0 R] >>".encode(),
            b"<< /S /JavaScript /JS (app.alert\\(1\\);) >>",
        ],
    )


def open_action_javascript_pdf() -> bytes:
    """JavaScript that runs automatically on open."""
    n = first_free_object()
    return build_pdf(
        catalog_extra=f"/OpenAction {n} 0 R",
        extra_objects=[b"<< /S /JavaScript /JS (app.alert\\(1\\);) >>"],
    )


def launch_action_pdf() -> bytes:
    """An annotation that launches an external program."""
    n = first_free_object()
    return build_pdf(
        page_extra=f"/Annots [{n} 0 R]",
        extra_objects=[
            f"<< /Type /Annot /Subtype /Link /Rect [0 0 100 100] "
            f"/A {n + 1} 0 R >>".encode(),
            b"<< /S /Launch /F (cmd.exe) >>",
        ],
    )


def embedded_file_pdf(filename: str = "payload.exe") -> bytes:
    """An embedded file attachment."""
    n = first_free_object()
    return build_pdf(
        catalog_extra=f"/Names << /EmbeddedFiles {n} 0 R >>",
        extra_objects=[
            f"<< /Names [({filename}) {n + 1} 0 R] >>".encode(),
            f"<< /Type /Filespec /F ({filename}) /EF << /F {n + 2} 0 R >> >>".encode(),
            b"<< /Length 5 /Type /EmbeddedFile >>\nstream\nMZ\x90\x00\nendstream",
        ],
    )


def suspicious_url_pdf(url: str = "http://192.168.1.10/payload.exe") -> bytes:
    """A link annotation pointing at a suspicious URL."""
    n = first_free_object()
    return build_pdf(
        page_extra=f"/Annots [{n} 0 R]",
        extra_objects=[
            f"<< /Type /Annot /Subtype /Link /Rect [0 0 100 100] "
            f"/A {n + 1} 0 R >>".encode(),
            f"<< /S /URI /URI ({url}) >>".encode(),
        ],
    )


def benign_url_pdf(url: str = "https://example.com/report") -> bytes:
    n = first_free_object()
    return build_pdf(
        page_extra=f"/Annots [{n} 0 R]",
        extra_objects=[
            f"<< /Type /Annot /Subtype /Link /Rect [0 0 100 100] "
            f"/A {n + 1} 0 R >>".encode(),
            f"<< /S /URI /URI ({url}) >>".encode(),
        ],
    )


def signature_field_pdf() -> bytes:
    """An AcroForm carrying a signature field."""
    n = first_free_object()
    return build_pdf(
        catalog_extra=f"/AcroForm << /Fields [{n} 0 R] >>",
        extra_objects=[b"<< /FT /Sig /T (Signature1) /Ff 0 >>"],
    )


def xfa_form_pdf() -> bytes:
    n = first_free_object()
    return build_pdf(
        catalog_extra=f"/AcroForm << /Fields [] /XFA {n} 0 R >>",
        extra_objects=[b"<< /Length 9 >>\nstream\n<xdp:xdp>\nendstream"],
    )


def additional_actions_pdf() -> bytes:
    """Document-level /AA additional actions."""
    n = first_free_object()
    return build_pdf(
        catalog_extra=f"/AA << /WC {n} 0 R >>",
        extra_objects=[b"<< /S /JavaScript /JS (close\\(\\);) >>"],
    )


def page_additional_actions_pdf() -> bytes:
    n = first_free_object()
    return build_pdf(
        page_extra=f"/AA << /O {n} 0 R >>",
        extra_objects=[b"<< /S /Movie >>"],
    )


def remote_goto_pdf() -> bytes:
    n = first_free_object()
    return build_pdf(
        page_extra=f"/Annots [{n} 0 R]",
        extra_objects=[
            f"<< /Type /Annot /Subtype /Link /Rect [0 0 10 10] /A {n + 1} 0 R >>".encode(),
            b"<< /S /GoToR /F (http://example.org/other.pdf) >>",
        ],
    )


def corrupt_pdf() -> bytes:
    """Truncated garbage — must be handled, not crash."""
    return b"%PDF-1.7\n1 0 obj\n<< /Type /Catalog"


def empty_text_pdf() -> bytes:
    """A page with no text operators, standing in for a scan."""
    return build_pdf(content=b"q 1 0 0 1 0 0 cm Q")


def fillable_form_pdf() -> bytes:
    """A real AcroForm built with reportlab: text, required text, multiline,
    checkbox, dropdown and a read-only field."""
    from reportlab.pdfgen import canvas
    from reportlab.lib.colors import black, white

    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=(612, 792))
    form = pdf.acroForm

    pdf.drawString(72, 740, "Application Form")

    pdf.drawString(72, 700, "Full name:")
    form.textfield(name="full_name", x=160, y=694, width=300, height=20,
                   borderColor=black, fillColor=white, textColor=black,
                   forceBorder=True, tooltip="Your legal name", maxlen=60)

    pdf.drawString(72, 660, "Email (required):")
    form.textfield(name="email", x=190, y=654, width=270, height=20,
                   borderColor=black, fillColor=white, textColor=black,
                   forceBorder=True, fieldFlags="required")

    pdf.drawString(72, 610, "Comments:")
    form.textfield(name="comments", x=160, y=570, width=300, height=60,
                   borderColor=black, fillColor=white, textColor=black,
                   forceBorder=True, fieldFlags="multiline")

    pdf.drawString(72, 530, "Agree to terms:")
    form.checkbox(name="agree", x=190, y=526, size=16,
                  borderColor=black, fillColor=white, forceBorder=True)

    pdf.drawString(72, 490, "Country:")
    form.choice(name="country", x=160, y=484, width=160, height=20,
                options=[("United Kingdom", "UK"), ("United States", "US"),
                         ("Germany", "DE")],
                borderColor=black, fillColor=white, textColor=black,
                forceBorder=True, value="UK")

    pdf.drawString(72, 450, "Reference:")
    form.textfield(name="reference", x=160, y=444, width=160, height=20,
                   borderColor=black, fillColor=white, textColor=black,
                   forceBorder=True, value="REF-001", fieldFlags="readOnly")

    pdf.save()
    return buffer.getvalue()


def as_stream(data: bytes, name: str = "corpus.pdf"):
    """Wrap bytes the way Streamlit hands an upload to the app."""
    stream = io.BytesIO(data)
    stream.name = name
    stream.size = len(data)
    return stream


# --- non-PDF fixtures, for the formats the uploader now converts ------------

DOCX_SENTENCE = "Revenue rose twelve percent this quarter."


def small_docx() -> bytes:
    """A minimal Word document containing DOCX_SENTENCE."""
    import io

    from docx import Document

    document = Document()
    document.add_heading("Quarterly review", 0)
    document.add_paragraph(DOCX_SENTENCE)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def small_xlsx() -> bytes:
    """A one-sheet workbook with a header row and one row of data."""
    import io

    from openpyxl import Workbook

    book = Workbook()
    sheet = book.active
    sheet["A1"], sheet["B1"] = "Region", "Total"
    sheet["A2"], sheet["B2"] = "North", 128400
    buffer = io.BytesIO()
    book.save(buffer)
    return buffer.getvalue()


def small_png(size: tuple = (400, 300)) -> bytes:
    import io

    from PIL import Image

    buffer = io.BytesIO()
    Image.new("RGB", size, "white").save(buffer, "PNG")
    return buffer.getvalue()
