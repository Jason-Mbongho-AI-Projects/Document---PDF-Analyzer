"""
Conversion.

A converter is not verified by producing bytes. Every test here opens the
output with the library that owns that format — python-docx, openpyxl, csv,
json, Pillow, pypdf — and asserts the content survived. A file that cannot be
opened is a failed conversion no matter what the endpoint returned.
"""
import csv as csv_module
import io
import json
import zipfile

import pytest
from PIL import Image
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Table, TableStyle

import pdf_corpus as corpus
from docintel.pdf import convert as C
from docintel.pdf.engine import PDFEngineError


def table_pdf() -> bytes:
    buffer = io.BytesIO()
    document = SimpleDocTemplate(buffer)
    styles = getSampleStyleSheet()
    rows = [
        ["Region", "Revenue", "Growth"],
        ["EMEA", "1200", "12"],
        ["APAC", "980", "8"],
        ["AMER", "1500", "15"],
    ]
    table = Table(rows)
    table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.black)]))
    document.build([Paragraph("Quarterly Results", styles["Title"]), table])
    return buffer.getvalue()


TABLE_PDF = table_pdf()


# ------------------------------------------------------------ capabilities

def test_capabilities_lists_targets_with_fidelity():
    caps = {c.target: c for c in C.capabilities()}
    assert caps["txt"].fidelity == "text-only"
    assert caps["xlsx"].fidelity == "structural"
    assert caps["png"].fidelity == "raster"


def test_unavailable_targets_explain_why():
    """Anything needing an office engine must say so, not silently vanish."""
    caps = {c.target: c for c in C.capabilities()}
    pptx = caps["pptx"]
    if not pptx.available:
        assert "LibreOffice" in pptx.reason


def test_unknown_target_is_rejected_with_the_list():
    with pytest.raises(PDFEngineError, match="Available targets"):
        C.convert(TABLE_PDF, "wordperfect")


# -------------------------------------------------- outputs actually open

def test_txt_contains_the_text():
    result = C.convert(TABLE_PDF, "txt", filename="results.pdf")
    body = result.data.decode("utf-8")
    assert "Quarterly Results" in body
    assert result.filename == "results.txt"


def test_markdown_is_produced_and_flags_its_heuristic():
    result = C.convert(TABLE_PDF, "markdown")
    assert b"Quarterly Results" in result.data
    assert any("heuristic" in w for w in result.warnings)


def test_html_is_well_formed_and_escapes_content():
    from reportlab.pdfgen import canvas
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=(612, 792))
    pdf.setFont("Helvetica", 12)
    pdf.drawString(72, 720, "Angle brackets <script> and & ampersand")
    pdf.save()

    body = C.convert(buffer.getvalue(), "html").data.decode()
    assert body.startswith("<!doctype html>")
    assert "&lt;script&gt;" in body
    assert "<script>" not in body


def test_json_parses_and_carries_geometry_and_tables():
    payload = json.loads(C.convert(TABLE_PDF, "json").data)
    assert payload["page_count"] == 1

    page = payload["pages"][0]
    assert page["words"], "no word geometry"
    assert set(page["words"][0]["view_rect"]) == {"x", "y", "width", "height"}
    assert page["tables"], "no tables detected"


def test_csv_parses_with_the_real_rows():
    text = C.convert(TABLE_PDF, "csv").data.decode()
    rows = [r for r in csv_module.reader(io.StringIO(text)) if r]

    assert ["Region", "Revenue", "Growth"] in rows
    assert ["EMEA", "1200", "12"] in rows


def test_xlsx_opens_in_openpyxl_with_numeric_cells():
    from openpyxl import load_workbook

    workbook = load_workbook(io.BytesIO(C.convert(TABLE_PDF, "xlsx").data))
    sheet = workbook[workbook.sheetnames[0]]
    rows = [[cell.value for cell in row] for row in sheet.iter_rows()]

    assert rows[0] == ["Region", "Revenue", "Growth"]
    # Numbers written as numbers, so the sheet is usable for arithmetic.
    assert rows[1][1] == 1200
    assert isinstance(rows[1][1], int)


def test_docx_opens_in_python_docx_with_text_and_tables():
    from docx import Document

    document = Document(io.BytesIO(C.convert(TABLE_PDF, "docx").data))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    assert any("Quarterly Results" in p for p in paragraphs)
    assert document.tables
    header = [cell.text for cell in document.tables[0].rows[0].cells]
    assert header == ["Region", "Revenue", "Growth"]


def test_docx_states_that_layout_is_not_preserved():
    result = C.convert(TABLE_PDF, "docx")
    assert result.fidelity == "text-only"
    assert any("layout" in w.lower() for w in result.warnings)


def test_images_zip_contains_one_valid_image_per_page():
    result = C.convert(corpus.multipage_pdf(3), "png")
    archive = zipfile.ZipFile(io.BytesIO(result.data))

    assert len(archive.namelist()) == 3
    for name in archive.namelist():
        image = Image.open(io.BytesIO(archive.read(name)))
        assert image.format == "PNG"
        assert image.width > 100


def test_jpeg_export_produces_jpegs():
    result = C.convert(corpus.clean_pdf(), "jpg")
    archive = zipfile.ZipFile(io.BytesIO(result.data))
    image = Image.open(io.BytesIO(archive.read(archive.namelist()[0])))
    assert image.format == "JPEG"


def test_csv_export_without_tables_says_so_instead_of_returning_junk():
    with pytest.raises(PDFEngineError, match="No tables were detected"):
        C.convert(corpus.clean_pdf(), "csv")


# ------------------------------------------------------------ to PDF

def test_text_to_pdf_opens_and_contains_the_text():
    output = C.to_pdf(b"Hello world.\n\nSecond paragraph.", "txt", title="Notes")
    reader = PdfReader(io.BytesIO(output))
    text = reader.pages[0].extract_text()

    assert "Hello world." in text
    assert "Notes" in text


def test_text_to_pdf_escapes_markup_rather_than_rendering_it():
    """User text must never be interpreted as reportlab markup."""
    output = C.to_pdf(b"A <b>bold</b> claim & more", "txt")
    text = PdfReader(io.BytesIO(output)).pages[0].extract_text()
    assert "<b>" in text        # shown literally, not applied


def test_csv_to_pdf_renders_a_table():
    output = C.to_pdf(b"Name,Amount\nAlice,100\nBob,200\n", "csv", title="Ledger")
    text = PdfReader(io.BytesIO(output)).pages[0].extract_text()

    assert "Alice" in text and "200" in text


def test_image_to_pdf_produces_a_page():
    image = Image.new("RGB", (400, 300), (30, 90, 200))
    buffer = io.BytesIO()
    image.save(buffer, "PNG")

    output = C.to_pdf(buffer.getvalue(), "png")
    reader = PdfReader(io.BytesIO(output))
    assert len(reader.pages) == 1
    assert float(reader.pages[0].mediabox.width) == pytest.approx(400, abs=1)


def test_unsupported_source_is_rejected():
    with pytest.raises(PDFEngineError, match="Cannot create a PDF"):
        C.to_pdf(b"data", "exe")


def test_office_to_pdf_reports_unavailability_honestly():
    if C._office_binary() is not None:
        pytest.skip("LibreOffice is installed on this machine")

    with pytest.raises(PDFEngineError, match="LibreOffice"):
        C.to_pdf(b"fake docx", "docx")


def test_empty_csv_is_rejected():
    with pytest.raises(PDFEngineError, match="empty"):
        C.to_pdf(b"", "csv")


# ------------------------------------------------------------------- api

def test_capabilities_endpoint(alice):
    body = alice.get("/api/v1/convert/capabilities").json()
    targets = {c["target"] for c in body["from_pdf"]}
    assert {"txt", "docx", "xlsx", "png"} <= targets
    assert "csv" in body["to_pdf"]
    assert "does not record paragraphs" in body["note"]


def test_convert_endpoint_returns_a_usable_file(alice):
    document_id = alice.upload(TABLE_PDF, name="results.pdf").json()["document"]["id"]

    response = alice.post(f"/api/v1/documents/{document_id}/convert",
                          json={"target": "xlsx"})
    assert response.status_code == 200
    assert "spreadsheetml" in response.headers["content-type"]
    assert response.headers["x-conversion-fidelity"] == "structural"

    from openpyxl import load_workbook
    workbook = load_workbook(io.BytesIO(response.content))
    assert workbook.sheetnames


def test_convert_endpoint_reports_fidelity_in_headers(alice):
    document_id = alice.upload(TABLE_PDF, name="r.pdf").json()["document"]["id"]
    response = alice.post(f"/api/v1/documents/{document_id}/convert",
                          json={"target": "docx"})
    assert response.headers["x-conversion-fidelity"] == "text-only"
    assert "layout" in response.headers["x-conversion-note"].lower() or \
           "Layout" in response.headers["x-conversion-warnings"]


def test_convert_leaves_the_source_untouched(alice):
    document_id = alice.upload(TABLE_PDF, name="r.pdf").json()["document"]["id"]
    before = alice.get(f"/api/v1/documents/{document_id}/download").content

    alice.post(f"/api/v1/documents/{document_id}/convert", json={"target": "txt"})

    after = alice.get(f"/api/v1/documents/{document_id}/download").content
    assert before == after
    assert alice.get(f"/api/v1/documents/{document_id}").json()["version_count"] == 1


def test_convert_bad_target_returns_400(alice):
    document_id = alice.upload(TABLE_PDF, name="r.pdf").json()["document"]["id"]
    response = alice.post(f"/api/v1/documents/{document_id}/convert",
                          json={"target": "nonsense"})
    assert response.status_code == 400
    assert "Available targets" in response.json()["detail"]


def test_to_pdf_endpoint_creates_a_document(alice):
    response = alice.client.post(
        "/api/v1/convert/to-pdf",
        headers=alice.headers,
        data={"workspace_id": alice.workspace_id},
        files={"file": ("notes.txt", b"Line one.\n\nLine two.", "text/plain")},
    )
    assert response.status_code == 201

    body = response.json()
    assert body["filename"] == "notes.pdf"
    assert body["converted_from"] == "txt"

    downloaded = alice.get(f"/api/v1/documents/{body['document_id']}/download").content
    assert "Line one." in PdfReader(io.BytesIO(downloaded)).pages[0].extract_text()


def test_to_pdf_rejects_an_unsupported_type(alice):
    response = alice.client.post(
        "/api/v1/convert/to-pdf",
        headers=alice.headers,
        data={"workspace_id": alice.workspace_id},
        files={"file": ("thing.exe", b"MZ\x90", "application/octet-stream")},
    )
    assert response.status_code == 400


def test_other_tenant_cannot_convert(alice, bob):
    document_id = alice.upload(TABLE_PDF, name="r.pdf").json()["document"]["id"]
    assert bob.post(f"/api/v1/documents/{document_id}/convert",
                    json={"target": "txt"}).status_code == 404
